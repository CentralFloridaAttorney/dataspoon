import docx


class DocxTool:
    def __init__(self):
        print('__init__ done!')

    @staticmethod
    def get_text(_file_path):
        doc = docx.Document(_file_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    @staticmethod
    def get_sentence_list(_file_path):
        doc = docx.Document(_file_path)
        text_list = []
        for para in doc.paragraphs:
            text_list.append(para.text)
        return text_list
